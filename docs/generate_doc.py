#!/usr/bin/env python3
"""Generate PCAS-Connect project documentation as a DOCX file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# --- Style Setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for s in doc.styles:
    if s.type is not None and hasattr(s, 'font'):
        s.font.name = 'Times New Roman'

def add_heading_custom(text, level=1, align=WD_ALIGN_PARAGRAPH.CENTER):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
                r.font.name = 'Times New Roman'
    for row_data in rows:
        row_cells = t.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
    return t

def page_break():
    doc.add_page_break()

# =============================================
# COVER PAGE
# =============================================
for _ in range(3):
    doc.add_paragraph()

add_para("PCAS-Connect", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22)
add_para("College Attendance and Academic Management System", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)

doc.add_paragraph()

add_para("Submitted in partial fulfilment of the Requirements for the award of the degree of", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para("BACHELOR OF COMPUTER APPLICATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)

doc.add_paragraph()
add_para("Submitted By", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()
add_para("MOHAMMED NASIK K", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
doc.add_paragraph()
add_para("Reg.no. 230021079198", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_paragraph()
add_para("Under the guidance of", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()
add_para("Mrs. JOXY JOSEPH", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)

for _ in range(3):
    doc.add_paragraph()

add_para("DEPARTMENT OF COMPUTER APPLICATIONS\nPRESENTATION COLLEGE OF APPLIED SCIENCES,\nPUTHENVELIKARA, ERNAKULAM-683594\nMARCH 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

page_break()

# =============================================
# CERTIFICATE
# =============================================
add_para("PRESENTATION COLLEGE OF APPLIED SCIENCES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("[Affiliated to Mahatma Gandhi University, Kottayam] Ph: 0484-2485400, email:presentationcollege2002@gmail.com", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

add_heading_custom("CERTIFICATE", level=1)

add_para('This is to certify that the project entitled \u201cPCAS-Connect: College Attendance and Academic Management System\u201d is a Bonafide record of the work done by MOHAMMED NASIK K with Reg.no 230021079198 of the Department of Computer Application, Presentation College of Applied Sciences, Puthenvelikara in partial fulfilment of requirements for the award of Bachelor of Computer Applications by Mahatma Gandhi University, Kottayam during the year 2023-2026.')

doc.add_paragraph()
add_para("Head of the Department", align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()
add_para("Project Guide", align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()
add_para("Submitted for the university examination held on \u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026.", align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()
add_para("Examiners: 1.", align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("2.", align=WD_ALIGN_PARAGRAPH.LEFT)

page_break()

# =============================================
# DECLARATION
# =============================================
add_heading_custom("DECLARATION", level=1)

add_para('I hereby declare that the project entitled \u201cPCAS-Connect: College Attendance and Academic Management System\u201d submitted to Mahatma Gandhi University, Kottayam in partial fulfilment of the requirements for the award of the degree Bachelor of Computer Applications is a record of original work done by me during the period of study (2023-2026) under the guidance of Mrs JOXY JOSEPH, Department of Computer Applications.')

for _ in range(4):
    doc.add_paragraph()

add_para("Place:", align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Date:                                                                                     MOHAMMED NASIK K", align=WD_ALIGN_PARAGRAPH.LEFT)

page_break()

# =============================================
# ACKNOWLEDGEMENT
# =============================================
add_heading_custom("ACKNOWLEDGEMENT", level=1)

add_para("\tI use this opportunity to express my sincere thanks and deep sense of gratitude to all those who helped me in various ways for the completion of my project. I give all the glory and honour to GOD ALMIGHTY, who showered his providences and graces on the successful completion of the work.")

add_para("\tI sincerely thank my parents for supporting me in completing this project on time. It would have been impossible without their help and support. With great pleasure, I sincerely thank Mr. FEBY F A S, Principal, Presentation College of Applied Sciences, Puthenvelikara for his entire support and encouragement and never-ending support throughout the course in this prestigious institution.")

add_para("\tI sincerely thank Mrs. SEEMA M.A., Head of the Department of Computer Application, Presentation College of Applied Sciences, Puthenvelikara for her constant encouragement and never-ending support throughout this project.")

add_para("\tI express my deep and sincere gratitude to Mrs. JOXY JOSEPH, Department of Computer Application, Presentation College of Applied Sciences, Puthenvelikara, for her guidance, help, encouragement and never-ending support and guidance throughout the project.")

add_para("\tI would like to express my profound thanks to all the faculty members in my department for their valuable suggestions during the period of my project. Last but not least I am indebted to people who have directly or indirectly helped me for the successful completion of the project.")

page_break()

# =============================================
# ABSTRACT
# =============================================
add_heading_custom("ABSTRACT", level=1)

add_para("\tPCAS-Connect is a comprehensive College Attendance and Academic Management System designed for Presentation College of Applied Sciences. The system integrates a Django REST Framework backend, a React Native (Expo) mobile application, and a React-based Admin Web Portal to create a unified platform for managing academic activities. The system supports four distinct user roles: Admin, Head of Department (HOD), Teacher, and Student, each with role-specific functionalities and dashboards.")

add_para("\tThe platform enables period-wise attendance tracking with real-time percentage calculation, internal marks management with a digital entry and approval workflow, timetable scheduling, department and batch management, subject enrollment, student promotions, and college-wide announcements. Teachers can mark attendance per period, enter and submit internal marks, and view their timetables. HODs can oversee departmental activities, approve marks, manage teachers and students, and create announcements. Students can view their attendance statistics, internal marks, timetables, enrolled subjects, and personal profiles through the mobile application.")

add_para("\tThe system replaces traditional manual record-keeping with a digital, real-time, and transparent workflow, ensuring accuracy, accountability, and efficiency in college academic administration. The modular architecture allows easy scalability and future feature additions.")

page_break()

# =============================================
# CONTENTS
# =============================================
add_heading_custom("CONTENTS", level=1)

contents = [
    ("1.", "INTRODUCTION"),
    ("2.", "SYSTEM STUDY AND ANALYSIS"),
    ("", "2.1. Existing System"),
    ("", "2.2. Proposed System"),
    ("", "2.3. Feasibility Study"),
    ("", "    2.3.1 Operational Feasibility"),
    ("", "    2.3.2 Technical Feasibility"),
    ("", "    2.3.3 Financial and Economic Feasibility"),
    ("", "2.4. Modules and Module Description"),
    ("3.", "SYSTEM REQUIREMENT SPECIFICATION"),
    ("", "3.1 Hardware Requirements"),
    ("", "3.2 Software Requirements"),
    ("4.", "COST AND BENEFITS ANALYSIS"),
    ("", "4.1. Cost"),
    ("", "4.2. Benefits"),
    ("5.", "SYSTEM DESIGN"),
    ("", "5.1. Input Design"),
    ("", "5.2. Output Design"),
    ("", "5.3. Proposed Design"),
    ("", "    5.3.1. Data Flow Diagram"),
    ("", "    5.3.2. E-R Diagram"),
    ("", "5.4. Database Design"),
    ("6.", "SYSTEM TESTING AND IMPLEMENTATION"),
    ("7.", "SYSTEM MAINTENANCE"),
    ("8.", "CONCLUSION"),
    ("9.", "FUTURE SCOPE"),
    ("10.", "APPENDIX"),
    ("", "10.1. Screenshots"),
    ("", "10.2. Coding"),
    ("11.", "BIBLIOGRAPHY"),
]

for num, title in contents:
    p = doc.add_paragraph()
    if num:
        p.add_run(f"    {num} {title}").bold = True
    else:
        p.add_run(f"        {title}")
    p.paragraph_format.space_after = Pt(2)

page_break()

# =============================================
# 1. INTRODUCTION
# =============================================
add_heading_custom("INTRODUCTION", level=1)
add_heading_custom("1. INTRODUCTION", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("\tPCAS-Connect is a comprehensive College Attendance and Academic Management System developed specifically for Presentation College of Applied Sciences (PCAS), Puthenvelikara. The system is designed to digitize and streamline the core academic processes of the college, including attendance tracking, internal marks management, timetable scheduling, student and teacher administration, and institutional announcements. In most educational institutions, these processes are still handled using paper registers, spreadsheets, or disconnected software tools, leading to data inconsistencies, delayed reporting, lack of transparency, and increased administrative burden.")

add_para("\tThe system is built using a modern three-tier architecture. The backend API is developed using Django REST Framework (DRF) with Python, providing robust and secure RESTful endpoints for all operations. The mobile application is built using React Native with the Expo framework, enabling cross-platform compatibility on both Android and iOS devices. The administrative web portal is developed using React with Vite, offering a fast and responsive interface for college administrators and department heads.")

add_para("\tPCAS-Connect defines four distinct user roles: Admin, Head of Department (HOD), Teacher, and Student. Each role has a dedicated dashboard and set of functionalities tailored to their responsibilities. The Admin manages college-wide operations such as department creation, batch management, student and teacher onboarding, and report generation. The HOD oversees departmental activities including timetable management, teacher-subject assignments, mark approvals, student promotions, and departmental announcements. Teachers can mark period-wise attendance, enter internal marks, view timetables, and manage their profiles. Students can view their attendance statistics, internal marks, enrolled subjects, timetable, and personal information through the mobile application.")

add_para("\tThe system uses SQLite as the database during development, with the capability to scale to PostgreSQL or MySQL for production deployment. All data interactions are handled through secure API calls with token-based authentication, ensuring data integrity and user privacy. The modular and scalable architecture enables easy addition of new features and modules in the future.")

add_para("\tBy replacing manual record-keeping with an automated, real-time, and transparent digital system, PCAS-Connect aims to improve efficiency, accuracy, accountability, and communication across the entire academic workflow of the institution.")

page_break()

# =============================================
# 2. SYSTEM STUDY AND ANALYSIS
# =============================================
add_heading_custom("SYSTEM STUDY AND ANALYSIS", level=1)
add_heading_custom("2. SYSTEM STUDY AND ANALYSIS", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("2.1. Existing System", bold=True)
add_para("\tIn most colleges, academic management activities are handled through manual or semi-digital methods. Attendance is recorded using paper registers, which are prone to errors, tampering, and loss. Internal marks are maintained in spreadsheets or physical records, making consolidation and verification difficult. Timetable preparation and distribution rely on printed schedules that are hard to update when changes occur. Communication between departments, teachers, and students happens through notice boards, verbal announcements, or messaging apps, leading to information gaps and delays.")
add_para("\tStudent information management, including enrollment details, department assignments, and semester tracking, is typically handled through disconnected systems or manual files. There is no centralized platform for students to view their academic progress in real time. Teachers lack a digital tool to efficiently record attendance per period, manage multiple subjects, or track their own teaching schedules. Overall, the absence of a unified system results in poor coordination, duplicated effort, delayed decision-making, and reduced transparency across the institution.")

add_para("2.2. Proposed System", bold=True)
add_para("\tPCAS-Connect is proposed as a unified digital platform that centralizes all critical academic management functions into a single system accessible through web and mobile interfaces. The system provides role-based access for Admin, HOD, Teacher, and Student users, each with dedicated dashboards and functionalities.")
add_para("\tThe Admin can manage departments, batches, subjects, students, and teachers through the web portal, and generate comprehensive reports. HODs can create and manage timetables, assign teachers to subjects, approve internal marks, promote students across semesters, and post departmental announcements. Teachers can mark period-wise attendance for their assigned subjects, enter and submit internal marks for approval, and view their daily schedules. Students can access their attendance statistics (daily, monthly, and subject-wise), view internal marks, check timetables, see enrolled subjects, and manage their profiles through the mobile application.")
add_para("\tThe system ensures data consistency through a centralized database, provides real-time updates through API-driven communication, and maintains transparency through structured approval workflows (e.g., marks submission and approval). Overall, PCAS-Connect enhances productivity, accountability, and communication across the entire academic ecosystem.")

add_para("2.3. Feasibility Study", bold=True)

add_para("2.3.1. Operational Feasibility", bold=True)
add_para("\tThe system is operationally feasible because it directly addresses the daily challenges faced by college administrators, teachers, and students. The interface is intuitive and role-specific, with clear navigation and dashboards for each user type. Teachers can quickly mark attendance from their mobile devices, HODs can manage departmental operations from the web portal, and students can access their academic information anytime. The automated workflows for attendance tracking, marks management, and announcements reduce manual effort and improve operational efficiency.")

add_para("2.3.2. Technical Feasibility", bold=True)
add_para("\tThe system is technically feasible as it is built using well-established and widely supported technologies. The backend uses Django REST Framework (Python), which provides robust API development, built-in authentication, and ORM-based database management. The mobile application uses React Native with Expo, enabling cross-platform deployment on Android and iOS from a single codebase. The admin web portal uses React with Vite for fast development and optimized builds. SQLite serves as the development database, with seamless upgrade paths to PostgreSQL or MySQL for production. All technologies are open-source, well-documented, and actively maintained, ensuring long-term support and compatibility.")

add_para("2.3.3. Financial and Economic Feasibility", bold=True)
add_para("\tThe financial feasibility of the system is high because all core development tools and frameworks used are open-source and free of charge (Python, Django, React Native, React, SQLite). No expensive software licenses or specialized hardware are required. The system can be hosted on standard college infrastructure or affordable cloud platforms. Economically, the system delivers significant value by reducing manual paperwork, eliminating data entry errors, saving administrative time, improving reporting accuracy, and enhancing overall academic management. The long-term benefits of improved efficiency, transparency, and accountability far outweigh the minimal development and hosting costs.")

add_para("2.4. Modules and Module Description", bold=True)
add_para("\tThe project consists of four main modules based on user roles, each described below:")

add_para("1. Admin Module", bold=True)
add_para("\tThe Admin Module is the highest-level administrative interface, accessible through the web portal. The Admin has complete control over college-wide operations and data management.")
add_para("Functions of the Admin Module:")
admin_funcs = [
    "Login: Admin logs in securely using credentials.",
    "Dashboard: View college-wide statistics including total departments, students, teachers, and attendance summaries.",
    "Department Management: Add, edit, and manage academic departments with unique codes.",
    "Batch Management: Create and manage student batches (e.g., 2023-2026) with active/inactive status.",
    "Subject Management: Add subjects with details including code, semester, credits, type (Core, SEC, VAC, DSE, MDC, AEC), and department association.",
    "Student Management: Add students with complete profiles (name, register number, email, department, semester, DOB, phone, address, blood group, profile image). Search, filter, edit, and view student details.",
    "Teacher Management: Add teachers with profiles (name, email, department, phone, DOB, date of joining, gender, qualification, role). Search, filter, edit, and view teacher details.",
    "Timetable Management: View and manage department-wise timetable configurations.",
    "Attendance Monitoring: View class-wise and department-wise attendance reports.",
    "Reports: Generate detailed department-wise and student-wise reports with attendance and mark data, with export to Excel and PDF.",
    "Student Promotion: Promote students across semesters in bulk.",
]
for i, f in enumerate(admin_funcs, 1):
    add_para(f"\t{i}. {f}")

add_para("2. HOD Module", bold=True)
add_para("\tThe HOD (Head of Department) Module provides department-specific management capabilities accessible through both the web portal and mobile application.")
add_para("Functions of the HOD Module:")
hod_funcs = [
    "Login: HOD logs in using secure credentials.",
    "Dashboard: View department statistics including total students, teachers, subjects, and recent announcements.",
    "Timetable Editor: Create, edit, and manage semester-wise timetables by assigning subjects and teachers to specific periods and days.",
    "Teacher-Subject Assignment: Assign teachers to subjects within the department.",
    "Student List: View and manage all students in the department with filtering by semester.",
    "Teacher List: View all teachers in the department with their assigned subjects.",
    "Internal Marks Approval: Review and approve internal marks submitted by teachers, locking them for finalization.",
    "Attendance Overview: Monitor attendance statistics across the department.",
    "Announcements: Create and broadcast announcements to the entire college, specific departments, or targeted audiences (students/teachers).",
    "Student Promotion: Promote students from one semester to the next within the department.",
    "Profile Management: View and update personal profile information.",
    "Reports: Generate department-specific academic performance and attendance reports.",
]
for i, f in enumerate(hod_funcs, 1):
    add_para(f"\t{i}. {f}")

add_para("3. Teacher Module", bold=True)
add_para("\tThe Teacher Module is accessible through the mobile application and provides tools for daily academic activities.")
add_para("Functions of the Teacher Module:")
teacher_funcs = [
    "Login: Teacher logs in using secure credentials with forced password change on first login.",
    "Dashboard: View today's schedule, upcoming classes, and quick access to attendance and marks.",
    "Attendance Marking: Mark period-wise attendance for assigned subjects, with student lists populated automatically based on enrollment.",
    "Past Attendance: View and review previously recorded attendance records.",
    "Internal Marks Entry: Enter test scores (Test 1 and Test 2) for enrolled students, save as draft, and submit for HOD approval.",
    "Internal Marks Overview: View submitted and approved marks across all assigned subjects.",
    "Timetable: View personal weekly timetable showing assigned periods, subjects, and timings.",
    "Subjects: View all assigned subjects with details including semester, credits, and type.",
    "Announcements: View college-wide and department-specific announcements.",
    "Profile Management: View and update personal information including profile image.",
    "Change Password: Update login password securely.",
]
for i, f in enumerate(teacher_funcs, 1):
    add_para(f"\t{i}. {f}")

add_para("4. Student Module", bold=True)
add_para("\tThe Student Module is accessible through the mobile application and provides students with access to their academic information.")
add_para("Functions of the Student Module:")
student_funcs = [
    "Login: Student logs in using register number and password with forced password change on first login.",
    "Dashboard: View personal ID card, today's attendance summary (present/absent/percentage), and quick navigation to other sections.",
    "Attendance Statistics: View today's attendance with period-wise breakdown, monthly attendance with subject-wise analysis, and overall percentage calculation.",
    "Internal Marks: View test scores (Test 1 and Test 2) and total marks for all enrolled subjects.",
    "Timetable: View personal weekly timetable with subject and teacher details for each period.",
    "Subjects: View all enrolled subjects showing subject code, name, credits, semester, and type.",
    "Profile: View personal information including name, register number, department, semester, email, phone, DOB, address, and blood group.",
    "Change Password: Update login password securely.",
    "Help & Support: Access help documentation and support information.",
]
for i, f in enumerate(student_funcs, 1):
    add_para(f"\t{i}. {f}")

page_break()

# =============================================
# 3. SYSTEM REQUIREMENT SPECIFICATIONS
# =============================================
add_heading_custom("SYSTEM REQUIREMENT SPECIFICATIONS", level=1)
add_heading_custom("3. SYSTEM REQUIREMENT SPECIFICATIONS", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("3.1. Hardware Requirements:", bold=True)
hw_items = [
    "Processor / Motherboard: Intel Core i3 / i5 or equivalent",
    "RAM: Minimum 4GB (8GB recommended for smooth development and emulator operations)",
    "Storage: 256GB or above (SSD recommended for faster execution)",
    "Display: Standard 14\"/15.6\" HD or Full HD Monitor",
    "Mobile Device: Android or iOS smartphone with Expo Go app installed",
    "Network: Wi-Fi connectivity for mobile app communication with backend server",
]
for item in hw_items:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(item).font.size = Pt(12)

add_para("3.2. Software Requirements:", bold=True)
sw_items = [
    "Operating System: Windows / Linux / macOS",
    "Backend Language: Python 3.10+",
    "Backend Framework: Django 4.x with Django REST Framework",
    "Frontend Framework (Mobile): React Native with Expo",
    "Frontend Framework (Admin Web): React with Vite",
    "Database: SQLite (Development) / PostgreSQL or MySQL (Production)",
    "Web Browser: Google Chrome / Mozilla Firefox",
    "Mobile Platform: Android / iOS (via Expo Go)",
    "Version Control: Git with GitHub",
    "Code Editor: Visual Studio Code",
]
for item in sw_items:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(item).font.size = Pt(12)

# Technology descriptions
doc.add_paragraph()
add_para("Python", bold=True)
add_para("\tPython is a high-level, general-purpose programming language known for its clean syntax, readability, and versatility. It is widely used in web development, data science, automation, and machine learning. Python's extensive standard library and rich ecosystem of third-party packages make it ideal for rapid application development. Its interpreted nature allows quick prototyping and testing, while its object-oriented capabilities enable scalable and maintainable code architecture.")

add_para("Django REST Framework (DRF)", bold=True)
add_para("\tDjango REST Framework is a powerful and flexible toolkit built on top of Django for building Web APIs. It provides authentication policies including token-based authentication, serialization for complex data types, viewsets and routers for rapid API development, and browsable API interface for easy testing. DRF's built-in support for input validation, pagination, and permissions makes it ideal for building secure and scalable backend services for mobile and web applications.")

add_para("React Native (Expo)", bold=True)
add_para("\tReact Native is an open-source framework developed by Meta for building native mobile applications using JavaScript and React. The Expo framework extends React Native by providing a managed workflow with pre-configured build tools, over-the-air updates, and access to native device APIs without requiring native code compilation. This combination enables rapid cross-platform mobile development with a single codebase for both Android and iOS, significantly reducing development time and complexity.")

add_para("React (Vite)", bold=True)
add_para("\tReact is a widely used JavaScript library for building user interfaces, particularly single-page applications. Combined with Vite, a next-generation build tool, it provides extremely fast development server startup, hot module replacement, and optimized production builds. React's component-based architecture promotes reusability and maintainability, while Vite's build optimizations ensure excellent performance in production environments.")

add_para("SQLite", bold=True)
add_para("\tSQLite is a lightweight, serverless, self-contained relational database engine. It stores the entire database as a single file, requiring no separate server process or configuration. SQLite is fully supported by Django's ORM and provides ACID-compliant transactions. It is ideal for development and testing environments, with straightforward migration paths to enterprise databases like PostgreSQL or MySQL for production deployment.")

add_para("HTML, CSS and JavaScript", bold=True)
add_para("\tHTML (HyperText Markup Language) provides the structural foundation for web pages. CSS (Cascading Style Sheets) controls visual presentation, layout, responsive design, and animations. JavaScript adds interactivity, dynamic behavior, and client-side logic. Together, these technologies form the core of modern web development and are used extensively in both the React-based admin web portal and the React Native mobile application.")

page_break()

# =============================================
# 4. COST AND BENEFITS ANALYSIS
# =============================================
add_heading_custom("COST AND BENEFITS ANALYSIS", level=1)
add_heading_custom("4. COST AND BENEFIT ANALYSIS", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("4.1. Cost:", bold=True)

add_para("4.1.1. Development Costs:", bold=True)
dev_costs = [
    "Costs for assigning software developers, UI designers, and system testers for backend, frontend, and mobile development.",
    "Expenses for development tools such as code editors (VS Code), testing environments (Expo Go), and collaboration platforms (GitHub).",
    "Time investment in planning, designing models, building APIs, developing mobile screens, creating the admin portal, testing, and deployment.",
]
for c in dev_costs:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(c).font.size = Pt(12)

add_para("4.1.2. Operational Costs:", bold=True)
op_costs = [
    "Server or cloud hosting charges for deploying the Django backend and serving the admin web portal.",
    "Regular maintenance, security updates, bug fixes, and feature enhancements to keep the system functional and secure.",
    "Training sessions for teachers, HODs, and administrative staff to effectively use the system.",
    "Network infrastructure and backup storage costs for maintaining attendance records, marks data, and user profiles.",
]
for c in op_costs:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(c).font.size = Pt(12)

add_para("4.1.3. Administrative Costs:", bold=True)
admin_costs = [
    "Basic office resources such as system equipment, internet connectivity, and utility usage during development.",
    "Documentation, reporting, and oversight during system implementation and deployment.",
    "Technical support staff for monitoring system health and assisting users with queries.",
]
for c in admin_costs:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(c).font.size = Pt(12)

add_para("4.2. Benefits:", bold=True)

add_para("4.2.1. Improved Academic Management Efficiency:", bold=True)
b1 = [
    "Automated period-wise attendance tracking eliminates manual register maintenance and reduces errors.",
    "Digital internal marks entry with approval workflow ensures accuracy and prevents unauthorized modifications.",
    "Centralized timetable management with real-time visibility across all user roles minimizes scheduling conflicts.",
    "Streamlined student and teacher onboarding reduces administrative paperwork and processing time.",
]
for b in b1:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(b).font.size = Pt(12)

add_para("4.2.2. Cost Savings:", bold=True)
b2 = [
    "Reduces reliance on paper registers, printed timetables, and physical notice boards.",
    "Minimizes administrative overhead by integrating attendance, marks, timetable, and communication into one platform.",
    "Eliminates the need for separate attendance tracking software, mark consolidation tools, or communication apps.",
    "Open-source technology stack eliminates software licensing costs.",
]
for b in b2:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(b).font.size = Pt(12)

add_para("4.2.3. Transparency and Accountability:", bold=True)
b3 = [
    "Students can view attendance and marks in real time, promoting self-awareness and academic accountability.",
    "HODs can monitor departmental performance through detailed reports and attendance overviews.",
    "Structured marks approval workflow (Teacher submits → HOD approves) ensures verified and auditable records.",
    "Comprehensive reports enable data-driven decision-making for promotions, workload balancing, and resource allocation.",
]
for b in b3:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(b).font.size = Pt(12)

page_break()

# =============================================
# 5. SYSTEM DESIGN
# =============================================
add_heading_custom("SYSTEM DESIGN", level=1)
add_heading_custom("5. SYSTEM DESIGN", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("5.1. Input Design:", bold=True)
add_para("\tThe input design ensures that data entered into the system is accurate, validated, and structured for efficient processing. The primary inputs include:")

inputs = [
    ("User Authentication Data:", "Login credentials (username/register number and password) entered by Admin, HOD, Teacher, and Student users. Token-based authentication validates each session."),
    ("Department and Batch Data:", "Information entered by the Admin for creating departments (name, code) and batches (name, start year, end year, active status)."),
    ("Subject Data:", "Details including subject name, code, semester, credits, type (Core, SEC, VAC, DSE, MDC, AEC), and department association entered by the Admin."),
    ("Student Registration Data:", "Complete student profiles including name, register number, email, department, semester, date of birth, phone number, address, blood group, and profile image."),
    ("Teacher Registration Data:", "Teacher profiles including name, email, department, phone, date of birth, date of joining, gender, qualification, role (Teacher/HOD), and profile image."),
    ("Attendance Data:", "Period-wise attendance status (Present/Absent) for each student marked by teachers for specific subjects on specific dates."),
    ("Internal Marks Data:", "Test 1 and Test 2 scores entered by teachers for enrolled students in each subject, submitted for HOD approval."),
    ("Timetable Data:", "Period-subject-teacher assignments for each day of the week, organized by department and semester, configured by HODs."),
    ("Announcement Data:", "Title, content, target audience (All, Students, Teachers, Department-specific), and optional semester filter entered by HODs."),
]
for title, desc in inputs:
    add_para(f"\t\u25aa {title}", bold=True)
    add_para(f"\t{desc}")

add_para("5.2. Output Design:", bold=True)
add_para("\tThe system provides clear, structured, and role-specific outputs that support academic management and decision-making:")

outputs = [
    ("Student Dashboard:", "Displays personal ID card, today's attendance summary with period-wise breakdown (present count, absent count, percentage), and navigation to detailed views."),
    ("Attendance Reports:", "Subject-wise and period-wise attendance summaries with percentages, available at daily and monthly granularity for students, and department-wide overviews for HODs."),
    ("Internal Marks Display:", "Tabular presentation of Test 1 scores, Test 2 scores, and total marks for each enrolled subject, with status indicators (Draft/Submitted/Approved)."),
    ("Timetable Views:", "Weekly timetable grids showing period-wise subject and teacher assignments, displayed for students, teachers, and HODs."),
    ("Department Reports:", "Comprehensive reports with attendance and marks data, filterable by department and student, exportable to Excel and PDF formats."),
    ("Announcement Feed:", "Chronological list of announcements with title, content, sender, date, and audience scope."),
    ("User Profiles:", "Detailed personal information cards for students and teachers with profile images, contact details, and academic/professional information."),
]
for title, desc in outputs:
    add_para(f"\t\u25aa {title}", bold=True)
    add_para(f"\t{desc}")

add_para("5.3. Proposed Design:", bold=True)
add_para("\tThe system is designed using a three-tier modular architecture with clear separation between the presentation layer (React/React Native), business logic layer (Django REST Framework), and data layer (SQLite/PostgreSQL). The architecture follows RESTful API principles for communication between frontend and backend components.")
add_para("\tThe backend modules include:")
modules = [
    "User Authentication Module (handles login, token management, password change, and role-based access control)",
    "Department and Batch Management Module (manages academic departments and student batches)",
    "Subject and Enrollment Module (manages subjects, credits, types, and student-subject enrollment)",
    "Attendance Module (records period-wise attendance with teacher, subject, period, and date associations)",
    "Internal Marks Module (manages test scores with draft-submit-approve workflow)",
    "Timetable Module (manages day-period-subject-teacher mappings for each department and semester)",
    "Announcement Module (handles creation and distribution of college-wide and department-specific announcements)",
    "Student Promotion Module (manages semester-wise student progression)",
    "Reports Module (generates department-wise and student-wise analytical reports)",
]
for m in modules:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(m).font.size = Pt(12)

add_para("\tThe architecture is scalable, maintainable, and follows the MVC (Model-View-Controller) pattern through Django's MTV (Model-Template-View) framework. The API-driven design enables independent development and deployment of frontend and backend components.")

doc.add_paragraph()
add_para("5.3.1. Data Flow Diagram", bold=True)
add_para("\tA Data Flow Diagram (DFD) illustrates the flow of data, the processes that act on the data, and the data stores used. In the PCAS-Connect system, the DFD shows how data flows between Admin, HOD, Teacher, Student, and the system database through various processes including authentication, attendance marking, marks entry, timetable management, and report generation.")
add_para("\t[Data Flow Diagrams to be inserted here]", bold=True)

doc.add_paragraph()
add_para("5.3.2. Entity Relationship Diagram", bold=True)
add_para("\t[ER Diagram to be inserted here]", bold=True)

doc.add_paragraph()
add_para("5.4. Database Design", bold=True)
add_para("\tThe database consists of the following tables:")

# Department Table
add_para("Department Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier for each department"],
        ["name", "VARCHAR(100)", "NOT NULL", "Name of the department"],
        ["code", "VARCHAR(10)", "NOT NULL", "Short code for the department"],
    ]
)
doc.add_paragraph()

# Batch Table
add_para("Batch Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier for each batch"],
        ["name", "VARCHAR(50)", "NOT NULL", "Batch name (e.g., 2023-2026)"],
        ["start_year", "INT", "NOT NULL", "Starting year of the batch"],
        ["end_year", "INT", "NOT NULL", "Ending year of the batch"],
        ["is_active", "BOOLEAN", "DEFAULT True", "Whether the batch is currently active"],
    ]
)
doc.add_paragraph()

# Subject Table
add_para("Subject Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier for each subject"],
        ["name", "VARCHAR(200)", "NOT NULL", "Name of the subject"],
        ["code", "VARCHAR(20)", "NOT NULL", "Subject code"],
        ["semester", "INT", "NOT NULL", "Semester number"],
        ["credit", "INT", "NOT NULL", "Credit value of the subject"],
        ["subject_type", "VARCHAR(10)", "NOT NULL", "Type: CORE, SEC, VAC, DSE, MDC, AEC"],
        ["department_id", "INT", "FOREIGN KEY", "References the department"],
    ]
)
doc.add_paragraph()

# Student Table
add_para("Student Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier for each student"],
        ["user_id", "INT", "FOREIGN KEY, NULL", "References Django auth user for login"],
        ["name", "VARCHAR(100)", "NOT NULL", "Full name of the student"],
        ["register_number", "VARCHAR(20)", "UNIQUE, NOT NULL", "University register number"],
        ["email", "VARCHAR(254)", "UNIQUE, NOT NULL", "Email address"],
        ["department_id", "INT", "FOREIGN KEY", "References the department"],
        ["semester", "INT", "NOT NULL", "Current semester"],
        ["dob", "DATE", "NULL allowed", "Date of birth"],
        ["phone_number", "VARCHAR(20)", "NULL allowed", "Contact phone number"],
        ["address", "TEXT", "NULL allowed", "Residential address"],
        ["blood_group", "VARCHAR(10)", "NULL allowed", "Blood group"],
        ["profile_image", "VARCHAR(100)", "NULL allowed", "Path to profile image"],
        ["requires_password_change", "BOOLEAN", "DEFAULT True", "Force password change on first login"],
    ]
)
doc.add_paragraph()

# Teacher Table
add_para("Teacher Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier for each teacher"],
        ["user_id", "INT", "FOREIGN KEY, NULL", "References Django auth user for login"],
        ["name", "VARCHAR(100)", "NOT NULL", "Full name of the teacher"],
        ["email", "VARCHAR(254)", "UNIQUE, NOT NULL", "Email address"],
        ["department_id", "INT", "FOREIGN KEY", "References the department"],
        ["phone", "VARCHAR(20)", "NULL allowed", "Contact phone number"],
        ["dob", "DATE", "NULL allowed", "Date of birth"],
        ["date_of_joining", "DATE", "NULL allowed", "Date of joining the institution"],
        ["gender", "VARCHAR(10)", "NULL allowed", "Gender (Male/Female/Other)"],
        ["qualification", "VARCHAR(100)", "NULL allowed", "Academic qualification"],
        ["role", "VARCHAR(10)", "DEFAULT TEACHER", "Role: TEACHER or HOD"],
        ["profile_image", "VARCHAR(100)", "NULL allowed", "Path to profile image"],
        ["is_hod", "BOOLEAN", "DEFAULT False", "Whether teacher is HOD"],
        ["requires_password_change", "BOOLEAN", "DEFAULT True", "Force password change on first login"],
    ]
)
doc.add_paragraph()

# Attendance Table
add_para("Attendance Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier for attendance record"],
        ["student_id", "INT", "FOREIGN KEY", "References the student"],
        ["subject_id", "INT", "FOREIGN KEY", "References the subject"],
        ["teacher_id", "INT", "FOREIGN KEY", "References the teacher who marked"],
        ["period_id", "INT", "FOREIGN KEY", "References the period"],
        ["date", "DATE", "NOT NULL", "Date of attendance"],
        ["status", "CHAR(1)", "NOT NULL", "P for Present, A for Absent"],
    ]
)
doc.add_paragraph()

# Period Table
add_para("Period Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier for each period"],
        ["number", "INT", "NOT NULL", "Period number (1, 2, 3...)"],
        ["start_time", "TIME", "NOT NULL", "Start time of the period"],
        ["end_time", "TIME", "NOT NULL", "End time of the period"],
    ]
)
doc.add_paragraph()

# TimeTable Table
add_para("TimeTable Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier"],
        ["department_id", "INT", "FOREIGN KEY", "References the department"],
        ["semester", "INT", "NOT NULL", "Semester number"],
        ["day", "VARCHAR(3)", "NOT NULL", "Day code (MON, TUE, WED, THU, FRI, SAT)"],
        ["period_id", "INT", "FOREIGN KEY", "References the period"],
        ["subject_id", "INT", "FOREIGN KEY", "References the subject"],
        ["teacher_id", "INT", "FOREIGN KEY", "References the teacher"],
    ]
)
doc.add_paragraph()

# InternalMark Table
add_para("Internal Mark Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier"],
        ["student_id", "INT", "FOREIGN KEY", "References the student"],
        ["subject_id", "INT", "FOREIGN KEY", "References the subject"],
        ["test_1_scored", "DECIMAL(5,2)", "DEFAULT 0", "Score obtained in Test 1"],
        ["test_1_total", "DECIMAL(5,2)", "DEFAULT 50", "Maximum marks for Test 1"],
        ["test_2_scored", "DECIMAL(5,2)", "DEFAULT 0", "Score obtained in Test 2"],
        ["test_2_total", "DECIMAL(5,2)", "DEFAULT 50", "Maximum marks for Test 2"],
        ["total", "DECIMAL(5,2)", "DEFAULT 0", "Total internal mark"],
        ["is_draft", "BOOLEAN", "DEFAULT True", "Whether mark is in draft state"],
        ["is_submitted", "BOOLEAN", "DEFAULT False", "Whether teacher has submitted"],
        ["is_approved", "BOOLEAN", "DEFAULT False", "Whether HOD has approved"],
        ["updated_at", "DATETIME", "AUTO", "Last updated timestamp"],
    ]
)
doc.add_paragraph()

# Enrollment Table
add_para("Enrollment Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier"],
        ["student_id", "INT", "FOREIGN KEY", "References the student"],
        ["subject_id", "INT", "FOREIGN KEY", "References the subject"],
        ["date_enrolled", "DATETIME", "AUTO", "Date of enrollment"],
    ]
)
doc.add_paragraph()

# TeacherSubject Table
add_para("Teacher Subject Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier"],
        ["teacher_id", "INT", "FOREIGN KEY", "References the teacher"],
        ["subject_id", "INT", "FOREIGN KEY", "References the subject"],
    ]
)
doc.add_paragraph()

# Announcement Table
add_para("Announcement Table", bold=True)
add_table(
    ["Field Name", "Type", "Constraints", "Description"],
    [
        ["id", "INT (AUTO_INCREMENT)", "PRIMARY KEY", "Unique identifier"],
        ["title", "VARCHAR(200)", "NOT NULL", "Announcement title"],
        ["content", "TEXT", "NOT NULL", "Announcement body text"],
        ["date", "DATETIME", "AUTO", "Date and time of announcement"],
        ["sender_id", "INT", "FOREIGN KEY", "References the user who created it"],
        ["audience", "VARCHAR(10)", "NOT NULL", "Target: ALL, STUDENTS, TEACHERS, DEPT"],
        ["department_id", "INT", "FOREIGN KEY, NULL", "Department if audience is DEPT"],
        ["semester", "INT", "NULL allowed", "Semester filter if applicable"],
    ]
)

page_break()

# =============================================
# 6. SYSTEM TESTING AND IMPLEMENTATION
# =============================================
add_heading_custom("SYSTEM TESTING AND IMPLEMENTATION", level=1)
add_heading_custom("6. SYSTEM TESTING AND IMPLEMENTATION", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("System Testing", bold=True)
add_para("\tSystem Testing is a crucial phase in the software development lifecycle that involves testing an integrated system to verify that it meets specified requirements. This type of testing assesses the overall behaviour and performance of the system, ensuring that all components work together as intended.")

add_para("Key aspects of system testing include:", bold=True)
tests = [
    "Functional Testing: Verifies that the system performs the expected functions correctly. This includes testing user authentication across all roles, attendance marking and retrieval, marks entry and approval workflows, timetable CRUD operations, and report generation.",
    "API Testing: Validates all REST API endpoints for correct request/response handling, proper HTTP status codes, authentication enforcement, and data serialization. Tools like Django's test client and Postman are used for API validation.",
    "Integration Testing: Ensures seamless communication between the Django backend, React Native mobile app, and React admin web portal through API calls. Tests verify data consistency across all three platforms.",
    "Performance Testing: Assesses system behavior under various conditions, including concurrent attendance marking by multiple teachers, bulk student data loading, and report generation with large datasets.",
    "Security Testing: Evaluates the system's ability to protect data through token-based authentication, role-based access control, and input validation. Tests ensure unauthorized users cannot access restricted endpoints.",
    "Usability Testing: Ensures the mobile app and web portal are user-friendly with intuitive navigation, clear labels, responsive layouts, and appropriate error messages.",
    "Compatibility Testing: Verifies the mobile app works correctly on different Android and iOS devices, and the web portal renders properly across Chrome, Firefox, and Edge browsers.",
    "Regression Testing: Ensures that new feature additions or bug fixes do not negatively impact existing functionality across the system.",
]
for i, t in enumerate(tests, 1):
    add_para(f"\t{i}. {t}")

add_para("\tA systematic testing approach is followed after integration and before final deployment. Identified defects are tracked, resolved, and retested to ensure a high-quality, reliable product.")

doc.add_paragraph()
add_para("Implementation", bold=True)
add_para("\tThe PCAS-Connect system is implemented using a three-tier architecture with clear separation of concerns:")

add_para("\tThe backend is built using Django REST Framework (Python). The project follows Django's standard app-based structure with separate apps for core models, user authentication, student operations, and teacher operations. The core app defines all database models including Department, Batch, Subject, Student, Teacher, Attendance, TimeTable, InternalMark, Period, Enrollment, TeacherSubject, and Announcement. Each app contains its own views, serializers, and URL configurations for handling API requests.")

add_para("\tThe mobile application is built using React Native with the Expo framework. The app follows a screen-based architecture with separate screens for each functionality (49 screens total). Screens are organized by user role: Student screens (Dashboard, Login, Profile, Subjects, TimeTable), Teacher screens (Dashboard, Login, Attendance, Internal Marks, Timetable, Past Attendance, Profile, Subjects, Announcements), and HOD screens (Dashboard, Timetable Editor, Student List, Teacher List, Internal Marks, Attendance Overview, Announcements, Promotion, Reports, Profile). Reusable components such as SideNavMenu and IDCardPanel are shared across screens.")

add_para("\tThe admin web portal is built using React with Vite. It provides page-based routing for administrative functions including Dashboard, Departments, Students, Teachers, Attendance, Timetable, Syllabus, and Reports. HOD-specific pages are organized in a separate subdirectory. The portal communicates with the backend through API service modules.")

add_para("\tAuthentication is handled through Django's built-in authentication system with token-based API access. Users receive authentication tokens upon login, which are included in all subsequent API requests. The system enforces role-based access control, ensuring each user can only access endpoints and data permitted for their role. A forced password change mechanism ensures security for newly created accounts.")

add_para("\tThe database tables and relationships are managed through Django's ORM (Object-Relational Mapping), which handles migrations, queries, and data integrity automatically. Media files such as profile images are stored in the media directory and served through Django's media handling configuration.")

page_break()

# =============================================
# 7. SYSTEM MAINTENANCE
# =============================================
add_heading_custom("SYSTEM MAINTENANCE", level=1)
add_heading_custom("7. SYSTEM MAINTENANCE", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("\tSystem Maintenance involves the processes and activities required to keep a software system operational, efficient, and up-to-date after its initial deployment. This phase is crucial for ensuring the system continues to meet user needs and functions correctly over its lifespan.")

add_para("Key aspects of system maintenance include:", bold=True)

maint = [
    "Corrective Maintenance: This involves fixing any defects or bugs that are discovered in the system after deployment. This includes resolving API errors, fixing data inconsistencies in attendance or marks records, addressing mobile app crashes, and correcting UI rendering issues. Regular monitoring ensures the system remains reliable and operates as expected.",
    "Adaptive Maintenance: This type of maintenance is necessary when there are changes in the operating environment, such as updates to Python, Django, React Native, Expo SDK, or mobile OS versions. It also covers database upgrades (e.g., migrating from SQLite to PostgreSQL for production), server configuration changes, and adaptation to new university regulations or academic calendar modifications.",
    "Perfective Maintenance: This focuses on improving and enhancing the system based on user feedback and changing requirements. Potential improvements include adding new report formats, enhancing the mobile UI, implementing push notifications, adding biometric attendance integration, or introducing new modules such as exam scheduling or result management.",
    "Preventive Maintenance: This proactive approach aims to anticipate and prevent potential issues. It includes regular database backups, code optimization, security audits, dependency updates, performance monitoring, API documentation updates, and system health checks to ensure long-term reliability and performance.",
]
for i, m in enumerate(maint, 1):
    add_para(f"\t{i}. {m}")

add_para("\tOverall, system maintenance is an ongoing process that helps extend the life of the system, improve its performance, and adapt it to changing academic requirements and technological environments. It is essential for sustaining the effectiveness and efficiency of the PCAS-Connect system over time.")

page_break()

# =============================================
# 8. CONCLUSION
# =============================================
add_heading_custom("CONCLUSION", level=1)
add_heading_custom("8. CONCLUSION", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("\tThe PCAS-Connect system provides substantial benefits to Presentation College of Applied Sciences by digitizing and streamlining core academic management processes. By integrating attendance tracking, internal marks management, timetable scheduling, student and teacher administration, and institutional announcements into a single unified platform accessible through web and mobile interfaces, the system significantly improves academic coordination, operational efficiency, and institutional transparency.")

add_para("\tThe role-based architecture ensures that each user \u2014 Admin, HOD, Teacher, and Student \u2014 has access only to the functionalities relevant to their responsibilities, maintaining data security and operational clarity. The automated attendance tracking system with period-wise granularity eliminates manual register maintenance and provides real-time visibility into student attendance patterns. The digital internal marks workflow with draft-submit-approve stages ensures accuracy, accountability, and audit trails for academic assessments.")

add_para("\tThe mobile application empowers teachers and students to manage their academic activities conveniently from their smartphones, while the admin web portal provides administrators and HODs with comprehensive management tools and analytical reports. The modern technology stack \u2014 Django REST Framework, React Native, and React \u2014 ensures the system is robust, scalable, and maintainable.")

add_para("\tOverall, the PCAS-Connect system is a valuable investment for any educational institution that aims to enhance academic management efficiency, reduce manual workload, improve data accuracy, promote transparency, and foster better communication and accountability among all stakeholders.")

page_break()

# =============================================
# 9. FUTURE SCOPE
# =============================================
add_heading_custom("FUTURE SCOPE", level=1)
add_heading_custom("9. FUTURE SCOPE", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("\tThe PCAS-Connect system has strong potential for further enhancement and expansion in educational environments. In the future, the system can be improved by integrating the following advanced features:")

futures = [
    "Push Notifications: Implement real-time push notifications for attendance alerts, marks publication, announcement broadcasts, and deadline reminders using Firebase Cloud Messaging or similar services.",
    "Biometric and QR-Code Attendance: Integrate biometric fingerprint or QR-code based attendance systems for more secure and automated attendance verification, reducing possibilities of proxy attendance.",
    "Examination and Result Management: Add modules for managing examination schedules, hall ticket generation, grade calculation, result publication, and semester-wise grade cards (SGPA/CGPA calculation).",
    "Parent Portal: Develop a dedicated parent module that allows parents to monitor their ward's attendance, marks, and academic progress, with automatic alerts for low attendance or poor performance.",
    "AI-Powered Analytics: Implement machine learning models for predicting student performance based on attendance patterns and marks trends, identifying at-risk students, and providing personalized academic recommendations.",
    "Library and Fee Management: Extend the system to include library book tracking, fine management, fee payment tracking, and receipt generation.",
    "Multi-College Support: Evolve the platform to support multiple affiliated colleges under a single university umbrella, enabling centralized academic monitoring and inter-college data aggregation.",
    "Offline Mode: Implement offline capabilities in the mobile application for attendance marking in areas with poor network connectivity, with automatic synchronization when connectivity is restored.",
    "Advanced Reporting and Dashboards: Enhance reporting with interactive charts, trend analysis, comparative dashboards, and scheduled automated report delivery via email.",
    "Cloud Deployment: Deploy the production system on cloud platforms (AWS, Google Cloud, or Azure) with auto-scaling, CDN for media delivery, and automated CI/CD pipelines for seamless updates.",
]
for i, f in enumerate(futures, 1):
    add_para(f"\t{i}. {f}")

add_para("\tWith continuous development and adoption, PCAS-Connect can evolve into a comprehensive Academic ERP system, widely beneficial for small, medium, and large-scale educational institutions.")

page_break()

# =============================================
# 10. APPENDIX
# =============================================
add_heading_custom("APPENDIX", level=1)
add_heading_custom("10. APPENDIX", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("10.1 Screenshots", bold=True)
add_para("\t[Screenshots of the application to be inserted here]", bold=True)
add_para("\tRecommended screenshots to include:")
screenshots = [
    "Student Login Screen",
    "Student Dashboard with ID Card and Today's Attendance",
    "Student Attendance Statistics (Daily and Monthly views)",
    "Student Internal Marks View",
    "Student Timetable View",
    "Student Profile Screen",
    "Teacher Login Screen",
    "Teacher Dashboard",
    "Teacher Attendance Marking Screen",
    "Teacher Internal Marks Entry Screen",
    "Teacher Timetable View",
    "HOD Dashboard",
    "HOD Timetable Editor",
    "HOD Internal Marks Approval Screen",
    "HOD Announcement Creation Screen",
    "Admin Web Login Page",
    "Admin Dashboard",
    "Admin Department Management Page",
    "Admin Student Management Page",
    "Admin Teacher Management Page",
    "Admin Reports Page",
]
for i, s in enumerate(screenshots, 1):
    add_para(f"\t{i}. {s}")

doc.add_paragraph()
add_para("10.2 Coding", bold=True)
add_para("\t[Key source code files to be inserted here]", bold=True)
add_para("\tRecommended code sections to include:")
code_sections = [
    "Backend Models (core/models.py) - Database model definitions",
    "Student API Views (student/views.py) - Student data endpoints",
    "Teacher API Views (teacher/views.py) - Teacher operations endpoints",
    "Authentication Views (users/views.py) - Login and token management",
    "Mobile Student Dashboard Screen",
    "Mobile Teacher Attendance Screen",
    "Admin Web Dashboard Page",
]
for i, c in enumerate(code_sections, 1):
    add_para(f"\t{i}. {c}")

page_break()

# =============================================
# 11. BIBLIOGRAPHY
# =============================================
add_heading_custom("BIBLIOGRAPHY", level=1)

refs = [
    "Django for APIs: Build Web APIs with Python and Django \u2014 William S. Vincent, WelcomeToCode.",
    "Two Scoops of Django: Best Practices for Django \u2014 Daniel Greenfeld and Audrey Greenfeld, Two Scoops Press.",
    "React Native in Action \u2014 Nader Dabit, Manning Publications.",
    "Learning React: Modern Patterns for Developing React Apps \u2014 Alex Banks and Eve Porcello, O\u2019Reilly Media.",
    "JavaScript: The Definitive Guide \u2014 David Flanagan, O\u2019Reilly Media.",
    "Database System Concepts \u2014 Abraham Silberschatz, Henry F. Korth, S. Sudarshan, McGraw-Hill.",
    "Official Django Documentation \u2014 https://docs.djangoproject.com/",
    "Django REST Framework Documentation \u2014 https://www.django-rest-framework.org/",
    "React Native Documentation \u2014 https://reactnative.dev/docs/getting-started",
    "Expo Documentation \u2014 https://docs.expo.dev/",
    "React Documentation \u2014 https://react.dev/",
    "Vite Documentation \u2014 https://vitejs.dev/",
    "Python Official Documentation \u2014 https://docs.python.org/3/",
    "SQLite Documentation \u2014 https://www.sqlite.org/docs.html",
    "Mozilla Developer Network (MDN) \u2014 https://developer.mozilla.org",
    "Stack Overflow Community \u2014 https://stackoverflow.com",
]
for r in refs:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(r).font.size = Pt(12)

# Save
output_path = '/home/mohammednasikk/Desktop/PCAS-Connect/docs/PCAS-Connect(nasik).docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
