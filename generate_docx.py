from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_guide():
    doc = Document()
    
    # Title
    title = doc.add_heading('PCAS-Connect: Project Presentation Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This document serves as your complete guide to successfully presenting your project, '
        'PCAS-Connect, to your college guide/teacher. It covers the setup, execution, ideal '
        'presentation flow, and potential Viva questions.'
    )

    # Section 2: Project Overview
    doc.add_heading('2. Project Overview (For Your Viva Context)', level=1)
    doc.add_paragraph('Project Name: PCAS-Connect')
    p_arch = doc.add_paragraph('Architecture & Tech Stack:')
    doc.add_paragraph('Backend: Django REST Framework (Handles core API logic, database management, and authentication).', style='List Bullet')
    doc.add_paragraph('Admin Web Panel: React.js + Vite (A dedicated web dashboard for the college administration to manage broad system entities).', style='List Bullet')
    doc.add_paragraph('Mobile App: React Native + Expo (A unified cross-platform application for Students, Teachers, and HODs).', style='List Bullet')

    # Section 3: Setup Guide
    doc.add_heading('3. Presentation Day Setup (Step-by-Step)', level=1)
    doc.add_paragraph('Before walking in for your presentation, ensure your system is completely ready. DO NOT do this during the presentation; keep everything running in the background beforehand.')
    
    doc.add_heading('Step 1: Get the Backend Online', level=2)
    doc.add_paragraph('1. Open your terminal.', style='List Number')
    doc.add_paragraph('2. Navigate to the backend directory.', style='List Number')
    doc.add_paragraph('3. Activate the virtual environment (source venv/bin/activate).', style='List Number')
    doc.add_paragraph('4. Run the server using: python manage.py runserver 0.0.0.0:8000', style='List Number')
    doc.add_paragraph('Important: Make sure your PC and mobile device are connected to the same Wi-Fi network!', style='List Bullet')

    doc.add_heading('Step 2: Get the Admin Web Portal Ready', level=2)
    doc.add_paragraph('1. Open a new terminal tab.', style='List Number')
    doc.add_paragraph('2. Navigate to the admin-web directory.', style='List Number')
    doc.add_paragraph('3. Run the development server using: npm run dev', style='List Number')
    doc.add_paragraph('4. Keep the tab open in your browser so you can switch to it instantly.', style='List Number')

    doc.add_heading('Step 3: Connect the Mobile App', level=2)
    doc.add_paragraph('1. Verify your local IP address (using hostname -I or ipconfig).', style='List Number')
    doc.add_paragraph('2. Update your config.js in both the mobile app and admin web if your Wi-Fi IP has changed.', style='List Number')
    doc.add_paragraph('3. In a new terminal, go to mobile/PCASConnect.', style='List Number')
    doc.add_paragraph('4. Run: npx expo start', style='List Number')
    doc.add_paragraph('5. Scan the QR code with the Expo Go app on your phone and leave the app running.', style='List Number')

    # Section 4: The Presentation Flow
    doc.add_heading('4. How to Present (The Ideal Flow)', level=1)
    doc.add_paragraph('Follow this flow to give a structured, impressive presentation that covers all aspects of your system architecture.')

    doc.add_heading('Phase A: Introduce the Admin Dashboard (The Core)', level=2)
    doc.add_paragraph('Start with the Web Portal. Explain that this is where the college admin operates.', style='List Bullet')
    doc.add_paragraph('Show the Teacher and Student Management Pages.', style='List Bullet')
    doc.add_paragraph('Briefly show how easily Departments, Subjects, and Batches can be managed, adding or editing an entity as a live example.', style='List Bullet')

    doc.add_heading('Phase B: The Mobile App - Teacher View', level=2)
    doc.add_paragraph('Switch to your mobile phone (or an emulator if you are screen-sharing).', style='List Bullet')
    doc.add_paragraph('Log in using a Teacher account.', style='List Bullet')
    doc.add_paragraph('Demonstrate the \'Mark Attendance\' process. Show how a teacher selects a batch, subject, and marks specific students absent or present.', style='List Bullet')
    doc.add_paragraph('Show the interface they use to assign and view Internal Marks.', style='List Bullet')

    doc.add_heading('Phase C: The Mobile App - Student View', level=2)
    doc.add_paragraph('Log out, and instantly log in using a Student account.', style='List Bullet')
    doc.add_paragraph('Explain that Students have a different level of access.', style='List Bullet')
    doc.add_paragraph('Show their dashboard, demonstrating how real-time attendance percentage is displayed right after the teacher marks it.', style='List Bullet')
    
    doc.add_heading('Phase D: Conclude with the HOD View (Optional/Bonus)', level=2)
    doc.add_paragraph('If time permits, log in as an HOD to show their specialized dashboard overarching multiple teachers/batches.', style='List Bullet')

    # Section 5: Common Viva Questions
    doc.add_heading('5. Expected Viva Questions & Answers', level=1)
    
    doc.add_heading('Q1: Why did you use Django for the backend instead of Node.js/Express?', level=2)
    p_q1 = doc.add_paragraph('Answer: ')
    p_q1.add_run('Django provides incredible out-of-the-box features like an integrated ORM, administrative interfaces, and built-in SQLite handling. It naturally structured our application very well and sped up the development of building secure REST APIs.')
    
    doc.add_heading('Q2: Why React Native instead of Java/Kotlin/Swift?', level=2)
    p_q2 = doc.add_paragraph('Answer: ')
    p_q2.add_run('React Native allows us to develop cross-platform (iOS and Android) simultaneously using a single JavaScript/React codebase. It ensures high performance and lowers the maintenance overhead.')
    
    doc.add_heading('Q3: How do the different apps (Web and Mobile) communicate with the central database?', level=2)
    p_q3 = doc.add_paragraph('Answer: ')
    p_q3.add_run('Through RESTful APIs. The Django backend exposes endpoints. Both the React Web Admin and the React Native Mobile app make HTTP requests (GET, POST, PUT, DELETE) to these endpoints and consume JSON data.')

    doc.add_heading('Q4: What was the most challenging part of the development?', level=2)
    p_q4 = doc.add_paragraph('Answer: ')
    p_q4.add_run('State management on the mobile app, ensuring real-time syncing of attendance data (or mention managing your complex IP configurations between devices on the network).')

    # Conclusion
    doc.add_paragraph('\n')
    doc.add_paragraph('Good luck! Speak clearly, and confidently refer back to your role-based architecture if you get stuck. You\'ve got this.')

    # Save
    doc.save('/home/mohammednasikk/Desktop/PCAS-Connect/docs/PCAS_Connect_Presentation_Guide.docx')
    print("Document successfully created!")

if __name__ == "__main__":
    create_guide()
